#!/usr/bin/env python3
"""
Điền form Đề nghị Giải Ngân kiêm Khế ước Nhận nợ (KUNN) - Ngân hàng OCB
"""
import sys, argparse
from datetime import date
from docx import Document
from docx.oxml import OxmlElement

# ─── Chuyển số tiền sang chữ ──────────────────────────────────────────────
DONVI = ['', 'một', 'hai', 'ba', 'bốn', 'năm', 'sáu', 'bảy', 'tám', 'chín']

def _doc3(n, is_first=True):
    tram, chuc, dv = n // 100, (n % 100) // 10, n % 10
    r = []
    if tram:
        r.append(DONVI[tram] + ' trăm')
        if chuc == 0 and dv: r.append('lẻ')
    if chuc == 1:
        r.append('mười')
        if dv: r.append(DONVI[dv])
    elif chuc:
        r.append(DONVI[chuc] + ' mươi')
        if   dv == 1: r.append('mốt')
        elif dv == 5: r.append('lăm')
        elif dv:      r.append(DONVI[dv])
    elif dv and (tram or not is_first):
        r.append(DONVI[dv])
    elif dv:
        r.append(DONVI[dv])
    return ' '.join(r)

def so_thanh_chu(n):
    if n == 0: return 'Không đồng chẵn'
    groups, temp = [], n
    while temp:
        groups.append(temp % 1000)
        temp //= 1000
    HANGSO = ['', 'nghìn', 'triệu', 'tỷ']
    parts = []
    for i in range(len(groups)-1, -1, -1):
        g = groups[i]
        if g == 0: continue
        txt = _doc3(g, is_first=(i == len(groups)-1))
        if HANGSO[i]: txt += ' ' + HANGSO[i]
        parts.append(txt)
    s = ' '.join(parts).strip()
    return s[0].upper() + s[1:] + ' đồng chẵn'

def fmt_tien(n):
    return f"{n:,.0f}".replace(',', '.')

# ─── Helpers ──────────────────────────────────────────────────────────────
def _is_ellipsis_only(text):
    """Kiểm tra đoạn chỉ chứa dấu chấm lửng (có thể có ngoặc đơn đóng)."""
    stripped = text.strip().rstrip(')')
    return bool(stripped) and all(c in '…. \t' for c in stripped)

def set_para_text(para, new_text):
    """Thay toàn bộ text của paragraph, giữ run đầu tiên."""
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.add_run(new_text)

def replace_in_para(para, old, new):
    if old in para.text:
        set_para_text(para, para.text.replace(old, new))
        return True
    return False

def replace_all(doc, old, new):
    for para in doc.paragraphs:
        replace_in_para(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para, old, new)

def delete_para(para):
    """Xóa paragraph khỏi document."""
    p = para._element
    p.getparent().remove(p)

# ─── Điền form ────────────────────────────────────────────────────────────
def fill_form(template_path, output_path, so_tien, lai_suat, ngay_tra_lai, nam=None):
    doc  = Document(template_path)
    today = date.today()
    nam   = nam or today.year

    # 1. Năm hồ sơ
    replace_all(doc, '/2025/KUNN-OCB-DN', f'/{nam}/KUNN-OCB-DN')
    replace_all(doc, '/2026/KUNN-OCB-DN', f'/{nam}/KUNN-OCB-DN')  # phòng khi template đã 2026

    # 2. Ngày ký
    dd = f"{today.day:02d}"
    mm = f"{today.month:02d}"
    yy = str(today.year)
    replace_all(doc, 'ngày ……… tháng ……… năm ………', f'ngày {dd} tháng {mm} năm {yy}')
    replace_all(doc, 'ngày………..tháng……….năm……..', f'ngày {dd} tháng {mm} năm {yy}')

    # 3. Số tiền trong bảng (Row 0)
    so_tien_fmt = fmt_tien(so_tien)
    bang_chu    = so_thanh_chu(int(so_tien))

    for table in doc.tables:
        for row in table.rows:
            cell = row.cells[1] if len(row.cells) > 1 else None
            if not cell: continue
            texts = [p.text for p in cell.paragraphs]
            # Nhận diện ô "Số tiền nhận nợ"
            has_amount_field = any('……………………………….đồng' in t or 'đồng chẵn' in t.lower() for t in texts)
            if not has_amount_field:
                continue

            paras = list(cell.paragraphs)
            replaced_amount = False
            replaced_bangchu = False
            to_delete = []

            for para in paras:
                t = para.text
                if '……………………………….đồng' in t and not replaced_amount:
                    set_para_text(para, f'{so_tien_fmt} đồng')
                    replaced_amount = True
                elif 'Bằng chữ' in t and not replaced_bangchu:
                    set_para_text(para, f'(Bằng chữ: {bang_chu})')
                    replaced_bangchu = True
                elif _is_ellipsis_only(t):
                    to_delete.append(para)

            for p in to_delete:
                delete_para(p)

    # 4. Ngày trả lãi đầu tiên
    replace_all(doc, '…../……/……….', ngay_tra_lai)
    replace_all(doc, '….../……./………..', ngay_tra_lai)

    doc.save(output_path)
    return {
        'so_tien_so' : so_tien_fmt,
        'so_tien_chu': bang_chu,
        'ngay_ky'    : f'{dd}/{mm}/{yy}',
        'lai_suat'   : lai_suat,
        'ngay_tra_lai': ngay_tra_lai,
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('template')
    parser.add_argument('output')
    parser.add_argument('--so-tien',       type=float, required=True)
    parser.add_argument('--lai-suat',      type=float, required=True)
    parser.add_argument('--ngay-tra-lai',  required=True)
    parser.add_argument('--nam',           type=int, default=None)
    args = parser.parse_args()

    info = fill_form(args.template, args.output,
                     args.so_tien, args.lai_suat, args.ngay_tra_lai, args.nam)
    print(f"✅ Đã điền form KUNN:")
    print(f"   Số tiền  : {info['so_tien_so']} đồng")
    print(f"   Bằng chữ : {info['so_tien_chu']}")
    print(f"   Ngày ký  : {info['ngay_ky']}")
    print(f"   Ngày trả lãi đầu: {info['ngay_tra_lai']}")
    print(f"   💾 Lưu tại: {args.output}")

if __name__ == '__main__':
    sys.exit(main())
